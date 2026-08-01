"""Rewrite / Fabrication Agent — tailors resume content to a job description.

This agent mirrors the JD language while honoring the user's fabrication
_tolerance setting and preserving all original employment dates.
"""
import json
import logging
import re
from pathlib import Path

from docx import Document

from agents.jd_analyzer import JDAnalyzer
from models.job_models import JobListing
from utils.llm_client import llm_client

logger = logging.getLogger(__name__)


# Mapping of tolerance levels to LLM instruction text.
_FABRICATION_INSTRUCTIONS = {
    "none": (
        "Fabrication tolerance: NONE. Be strictly factual. Do not invent degrees, companies, "
        "tools, or responsibilities the candidate has not actually had. Only rephrase and "
        "emphasize existing experience."
    ),
    "moderate": (
        "Fabrication tolerance: MODERATE. You may rephrase and selectively emphasize existing "
        "skills and achievements to align with the JD. You may use project or team language "
        "that reflects real work, but do not invent degrees, companies, or tools the candidate "
        "has not used."
    ),
    "aggressive": (
        "Fabrication tolerance: AGGRESSIVE. Mirror the JD's exact terminology and rewrite "
        "bullets so the candidate's experience reads as a strong match. You may reframe role "
        "titles and project descriptions to align with the target role, but do not invent degrees, "
        "companies, dates, or tools the candidate has never used."
    ),
}


class ResumeTailor:
    """Generate tailored resume content from a base resume and a job listing."""

    def __init__(self, base_resume_path: Path) -> None:
        self.base_resume_path = base_resume_path
        self.base_resume_text = self._extract_text(base_resume_path)
        self.date_ranges = self._extract_date_ranges(base_resume_path)

    def _extract_text(self, path: Path) -> str:
        """Extract all text from the base resume docx."""
        doc = Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    tables.append(" | ".join(cells))
        return "\n\n".join(paragraphs + tables)

    def _extract_date_ranges(self, path: Path) -> list[str]:
        """Pull every employment date range out of the base resume.

        These ranges must be preserved unchanged in any generated resume so the
        candidate can verify exactly when each role was held.
        """
        doc = Document(str(path))
        ranges = []
        pattern = re.compile(
            r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s*\d{4}\s*" +
            r"[-–—]\s*(Present|Current|Now|Today|" +
            r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s*\d{4})",
            re.IGNORECASE,
        )
        for p in doc.paragraphs:
            for match in pattern.finditer(p.text):
                ranges.append(match.group(0))
        return ranges

    def tailor(self, job: JobListing, fabrication_tolerance: str = "moderate") -> dict:
        """Return tailored resume content as a structured dict."""
        job_analysis = JDAnalyzer().analyze(job)
        tolerance = fabrication_tolerance.lower()
        if tolerance not in _FABRICATION_INSTRUCTIONS:
            tolerance = "moderate"

        prompt = f"""You are an expert resume writer and ATS optimizer.

You will receive:
1. A base resume
2. A job description analysis
3. A fabrication tolerance level
4. The original employment date ranges that must be preserved exactly

Your task is to tailor the resume for the specific job while preserving factual accuracy and the candidate's real experience.

Return ONLY a JSON object with these exact keys:
- professional_title: A concise title line for the resume (e.g., "Data Analyst | SQL | Python | Power BI | Statistics | 6+ Years")
- professional_summary: 3-5 sentences summarizing the candidate as a strong fit for this role
- technical_skills: A list of skill categories matching the template. Each item is an object with {{category: str, skills: str}}. Use the same categories as the base resume if possible.
- experience: A list of work experiences. Each item is an object with {{job_header: str, bullets: [str, str, ...]}}. Keep the same job headers as the base resume where possible, but rewrite bullets to emphasize data analysis, SQL, reporting, visualization, and insights relevant to the target role. THE EMPLOYMENT DATE RANGES IN EACH job_header MUST MATCH ONE OF THE ORIGINAL DATE RANGES EXACTLY.

Base Resume:
{self.base_resume_text}

Job Analysis:
{json.dumps(job_analysis, indent=2)}

Original Employment Date Ranges (preserve exactly):
{chr(10).join(f"- {dr}" for dr in self.date_ranges)}

Fabrication Tolerance Instructions:
{_FABRICATION_INSTRUCTIONS[tolerance]}

Rules:
- Do not invent degrees, companies, or tools the candidate has not used.
- Emphasize data analysis, SQL, Python, Excel, Power BI/Tableau, reporting, dashboards, and insights.
- Use strong action verbs and quantify achievements where possible.
- Keep bullets concise (1-2 lines each).
- Keep the same overall resume structure and section headings.
- The professional title and summary should be specifically tailored to the target job title: {job.title}
"""
        messages = [
            {
                "role": "system",
                "content": "You are an expert resume writer. Return only valid JSON.",
            },
            {"role": "user", "content": prompt},
        ]
        raw = llm_client.chat(messages, temperature=0.4, max_tokens=4000)

        try:
            text = raw.strip()
            if text.startswith("```json"):
                text = text[7:]
            if text.startswith("```"):
                text = text[3:]
            if text.endswith("```"):
                text = text[:-3]
            return json.loads(text.strip())
        except json.JSONDecodeError as exc:
            logger.error("Failed to parse tailored resume JSON: %s\nRaw: %s", exc, raw)
            raise
