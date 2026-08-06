"""Build tailored DOCX and PDF resumes from the base template."""

import logging
import re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Try importing docx2pdf; fallback to a manual COM conversion if needed
try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except Exception:  # pragma: no cover
    DOCX2PDF_AVAILABLE = False

from config.settings import Settings

logger = logging.getLogger(__name__)


class ResumeBuilder:
    """Apply tailored content to the base resume template and export PDF."""

    def __init__(self, template_path: Path, output_dir: Path) -> None:
        self.template_path = template_path
        self.output_dir = output_dir

    def _find_heading_index(self, doc: Document, heading_text: str) -> int | None:
        """Return the paragraph index of the heading matching the text."""
        heading_text_lower = heading_text.lower()
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().lower() == heading_text_lower:
                return i
        return None

    def _replace_paragraph_text(self, paragraph, new_text: str) -> None:
        """Replace all runs in a paragraph while preserving formatting."""
        if not paragraph.runs:
            paragraph.add_run(new_text)
            return
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""

    def _replace_title(self, doc: Document, title: str) -> None:
        """Replace the professional title line (paragraph 1 in the template)."""
        if len(doc.paragraphs) > 1:
            self._replace_paragraph_text(doc.paragraphs[1], title)

    def _replace_summary(self, doc: Document, summary: str) -> None:
        """Replace the paragraph(s) under PROFESSIONAL SUMMARY."""
        idx = self._find_heading_index(doc, "PROFESSIONAL SUMMARY")
        if idx is None:
            logger.warning("PROFESSIONAL SUMMARY heading not found")
            return

        # Replace the first non-empty paragraph after the heading
        for i in range(idx + 1, len(doc.paragraphs)):
            p = doc.paragraphs[i]
            text = p.text.strip()
            if text and not text.isupper():
                self._replace_paragraph_text(p, summary)
                break

    def _replace_skills_table(self, doc: Document, skills: list[dict]) -> None:
        """Replace the technical skills table content."""
        idx = self._find_heading_index(doc, "TECHNICAL SKILLS")
        if idx is None:
            logger.warning("TECHNICAL SKILLS heading not found")
            return

        # Find the first table after the heading
        for table in doc.tables:
            if table.rows and len(table.rows[0].cells) == 2:
                # Clear existing rows
                for row in table.rows:
                    for cell in row.cells:
                        cell.text = ""
                # Populate with new skills
                for i, skill in enumerate(skills):
                    if i >= len(table.rows):
                        break
                    table.rows[i].cells[0].text = skill.get("category", "")
                    table.rows[i].cells[1].text = skill.get("skills", "")
                break

    def _is_job_header(self, text: str) -> bool:
        """Return True if the paragraph looks like a job header (contains a date)."""
        return bool(re.search(r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s*\d{4}", text))

    def _replace_experience(self, doc: Document, experience: list[dict]) -> None:
        """Replace experience bullets under each job header."""
        idx = self._find_heading_index(doc, "PROFESSIONAL EXPERIENCE")
        if idx is None:
            logger.warning("PROFESSIONAL EXPERIENCE heading not found")
            return

        # Find job header indices in the template
        header_indices = []
        for i in range(idx + 1, len(doc.paragraphs)):
            text = doc.paragraphs[i].text.strip()
            if self._is_job_header(text):
                header_indices.append(i)
            elif text.isupper() and text not in {"PROFESSIONAL EXPERIENCE"}:
                # Next section heading reached
                break

        # Replace job headers and bullets
        for job_idx, header_idx in enumerate(header_indices):
            if job_idx >= len(experience):
                break

            job = experience[job_idx]
            job_header = job.get("job_header", "")
            bullets = job.get("bullets", [])

            # Replace header
            if job_header:
                self._replace_paragraph_text(doc.paragraphs[header_idx], job_header)

            # Determine end of this job section
            next_header_idx = header_indices[job_idx + 1] if job_idx + 1 < len(header_indices) else None
            bullet_idx = 0

            for i in range(header_idx + 1, next_header_idx or len(doc.paragraphs)):
                if i >= len(doc.paragraphs):
                    break
                p = doc.paragraphs[i]
                text = p.text.strip()
                if not text:
                    continue
                if self._is_job_header(text) or (text.isupper() and len(text) > 3):
                    break
                if bullet_idx < len(bullets):
                    self._replace_paragraph_text(p, bullets[bullet_idx])
                    bullet_idx += 1

    def build(
        self, tailored_content: dict, job_title: str, company: str, job_id: str = ""
    ) -> tuple[Path, Path, str]:
        """Build a tailored DOCX and PDF resume.

        Returns (docx_path, pdf_path, base_name).
        """
        doc = Document(str(self.template_path))

        self._replace_title(doc, tailored_content.get("professional_title", ""))
        self._replace_summary(doc, tailored_content.get("professional_summary", ""))
        self._replace_skills_table(doc, tailored_content.get("technical_skills", []))
        self._replace_experience(doc, tailored_content.get("experience", []))

        today_str = datetime.now().strftime("%Y%m%d")
        safe_title = re.sub(r"[^\w\s-]", "", job_title).strip().replace(" ", "_")
        safe_company = re.sub(r"[^\w\s-]", "", company).strip().replace(" ", "_") or "Unknown"
        id_suffix = f"_{job_id}" if job_id else ""
        base_name = f"JD_{safe_title}_Lokesh{id_suffix}_{today_str}"

        docx_path = self.output_dir / f"{base_name}.docx"
        pdf_path = self.output_dir / f"{base_name}.pdf"

        doc.save(str(docx_path))
        logger.info("Saved tailored DOCX: %s", docx_path)

        self._convert_to_pdf(docx_path, pdf_path)
        logger.info("Saved tailored PDF: %s", pdf_path)

        return docx_path, pdf_path, base_name

    def _convert_to_pdf(self, docx_path: Path, pdf_path: Path) -> None:
        """Convert a DOCX file to PDF using the best available method."""
        if DOCX2PDF_AVAILABLE:
            try:
                convert(str(docx_path), str(pdf_path))
                return
            except Exception as exc:
                logger.warning("docx2pdf failed: %s", exc)

        # Fallback 1: Microsoft Word COM automation (Windows only).
        try:
            self._convert_with_word_com(docx_path, pdf_path)
            return
        except Exception as exc:
            logger.warning("Word COM conversion failed: %s", exc)

        # Fallback 2: LibreOffice headless conversion.
        try:
            self._convert_with_libreoffice(docx_path, pdf_path)
            return
        except Exception as exc:
            logger.warning("LibreOffice conversion failed: %s", exc)

        raise RuntimeError(
            "Could not convert DOCX to PDF. Ensure Microsoft Word or LibreOffice is available, "
            "or install docx2pdf."
        )

    def _convert_with_libreoffice(self, docx_path: Path, pdf_path: Path) -> None:
        """Convert DOCX to PDF using LibreOffice in headless mode."""
        import shutil
        import subprocess

        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice:
            raise RuntimeError("LibreOffice executable not found in PATH")

        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        cmd = [
            str(soffice),
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(docx_path),
        ]
        subprocess.run(cmd, check=True, capture_output=True, text=True, timeout=120)

        # LibreOffice names the output based on the input file name.
        expected = pdf_path.parent / f"{docx_path.stem}.pdf"
        if expected.exists() and expected.resolve() != pdf_path.resolve():
            expected.replace(pdf_path)

        if not pdf_path.exists():
            raise RuntimeError("LibreOffice finished but PDF was not created")

    def _convert_with_word_com(self, docx_path: Path, pdf_path: Path) -> None:
        """Convert DOCX to PDF using Microsoft Word COM."""
        import win32com.client as wc
        word = wc.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(str(docx_path))
            doc.SaveAs(str(pdf_path), FileFormat=17)  # 17 = PDF
            doc.Close()
        finally:
            word.Quit()


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    builder = ResumeBuilder(Settings.BASE_RESUME_TEMPLATE, Settings.OUTPUT_RESUME_DIR)
    sample_content = {
        "professional_title": "Data Analyst | SQL | Python | Power BI | Statistics | 6+ Years",
        "professional_summary": "Data Analyst with 6+ years of experience transforming complex datasets into actionable insights...",
        "technical_skills": [
            {"category": "Analytics & BI", "skills": "SQL, Python, Power BI, Tableau, Excel, Statistics, Data Visualization"},
            {"category": "Data Engineering", "skills": "ETL/ELT, Data Warehousing, Azure Databricks, Delta Lake, PySpark"},
            {"category": "Cloud & Tools", "skills": "Azure, AWS, Git, Docker, REST APIs, Pandas, NumPy"},
        ],
        "experience": [
            {
                "job_header": "Data Analyst | PepsiCo | Dec 2025 – Present | Dallas, TX",
                "bullets": [
                    "Analyzed large-scale operational datasets using SQL and Python to identify cost-saving opportunities.",
                    "Built interactive dashboards in Power BI to visualize KPIs for executive stakeholders.",
                ],
            }
        ],
    }
    docx, pdf = builder.build(sample_content, "Data Analyst", "State of Texas", job_id="sample123")
    print("DOCX:", docx)
    print("PDF:", pdf)
