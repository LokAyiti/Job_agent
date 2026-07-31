"""Generate tailored cover letters from a base template or from scratch."""

import logging
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except Exception:  # pragma: no cover
    DOCX2PDF_AVAILABLE = False

from utils.llm_client import llm_client
from models.job_models import JobListing

logger = logging.getLogger(__name__)


COVER_LETTER_PROMPT = """You are a professional cover letter writer.

Write a concise, personalized cover letter for the following job. The candidate is Lokesh Ayiti.

Return ONLY the cover letter body text (no date, no addresses, no signature block). The letter should be 3-4 paragraphs.

Job Title: {title}
Company: {company}
Location: {location}

Job Description:
{description}

Requirements:
{requirements}

Resume Highlights:
{resume_highlights}

Cover Letter Angle:
{angle}

Tone: professional, confident, concise. Mention specific skills from the resume that match the job.
"""


class CoverLetterBuilder:
    """Generate tailored cover letter DOCX and PDF files."""

    def __init__(self, output_dir: Path) -> None:
        self.output_dir = output_dir

    def generate_text(self, job: JobListing, resume_highlights: str, angle: str) -> str:
        """Generate cover letter text using the LLM."""
        prompt = COVER_LETTER_PROMPT.format(
            title=job.title,
            company=job.company or "Hiring Manager",
            location=job.location or "",
            description=job.description or "",
            requirements=job.requirements or "",
            resume_highlights=resume_highlights,
            angle=angle,
        )
        messages = [
            {
                "role": "system",
                "content": "You are a professional cover letter writer. Return only the cover letter body.",
            },
            {"role": "user", "content": prompt},
        ]
        return llm_client.chat(messages, temperature=0.4)

    def build(
        self, job: JobListing, resume_highlights: str, angle: str
    ) -> Path:
        """Build a tailored cover letter PDF.

        Returns the PDF path.
        """
        letter_text = self.generate_text(job, resume_highlights, angle)

        doc = Document()
        sections = doc.sections[0]
        sections.top_margin = Inches(1)
        sections.bottom_margin = Inches(1)
        sections.left_margin = Inches(1)
        sections.right_margin = Inches(1)

        # Header
        header = doc.add_paragraph()
        header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = header.add_run("Lokesh Ayiti\n")
        run.bold = True
        run.font.size = Pt(12)
        header.add_run("(815) 246-2333 | lokesh.1149a@gmail.com\n").font.size = Pt(10)
        header.add_run(f"{datetime.now().strftime('%B %d, %Y')}\n\n").font.size = Pt(10)

        # Recipient
        recipient = doc.add_paragraph()
        recipient.add_run(f"Hiring Manager\n{job.company or 'Company'}\n\n").font.size = Pt(11)

        # Body
        for para in letter_text.split("\n\n"):
            p = doc.add_paragraph()
            p.add_run(para.strip()).font.size = Pt(11)
            p.space_after = Pt(12)

        # Closing
        closing = doc.add_paragraph()
        closing.add_run("\n\nSincerely,\n\nLokesh Ayiti").font.size = Pt(11)

        today_str = datetime.now().strftime("%Y%m%d")
        safe_title = re.sub(r"[^\w\s-]", "", job.title).strip().replace(" ", "_")
        safe_company = re.sub(r"[^\w\s-]", "", job.company or "Unknown").strip().replace(" ", "_")
        id_suffix = f"_{job.job_id}" if job.job_id else ""
        base_name = f"CoverLetter_{safe_title}_{safe_company}_Lokesh{id_suffix}_{today_str}"

        docx_path = self.output_dir / f"{base_name}.docx"
        pdf_path = self.output_dir / f"{base_name}.pdf"

        doc.save(str(docx_path))
        logger.info("Saved cover letter DOCX: %s", docx_path)

        self._convert_to_pdf(docx_path, pdf_path)
        logger.info("Saved cover letter PDF: %s", pdf_path)

        return pdf_path

    def _convert_to_pdf(self, docx_path: Path, pdf_path: Path) -> None:
        """Convert DOCX to PDF."""
        if DOCX2PDF_AVAILABLE:
            try:
                convert(str(docx_path), str(pdf_path))
                return
            except Exception as exc:
                logger.warning("docx2pdf failed: %s", exc)

        try:
            self._convert_with_word_com(docx_path, pdf_path)
            return
        except Exception as exc:
            logger.warning("Word COM conversion failed: %s", exc)

        raise RuntimeError(
            "Could not convert cover letter DOCX to PDF. Ensure Microsoft Word is installed or LibreOffice is available."
        )

    def _convert_with_word_com(self, docx_path: Path, pdf_path: Path) -> None:
        """Convert DOCX to PDF using Microsoft Word COM."""
        import win32com.client as wc
        word = wc.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(str(docx_path))
            doc.SaveAs(str(pdf_path), FileFormat=17)  # 17 = PDF
            doc.Close()
        finally:
            word.Quit()


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    from config.settings import Settings
    sample = JobListing(
        title="Data Analyst",
        company="State of Texas",
        location="Austin, TX",
        description="Analyze public health data...",
        requirements="SQL, Python, 2 years experience...",
    )
    builder = CoverLetterBuilder(Settings.OUTPUT_COVER_LETTER_DIR)
    pdf = builder.build(sample, "SQL, Python, Power BI, data analysis", "Match public health analytics needs")
    print(pdf)
