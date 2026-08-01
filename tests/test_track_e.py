"""Tests for Track E — Resume Tailoring Engine (4 sub-agents + ledger)."""
import json
import sys
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

# job_application_system uses top-level relative imports.
_PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(_PROJECT_ROOT / "job_application_system"))

from agents.consistency_ledger import ConsistencyLedger
from models.job_models import JobListing


@pytest.fixture
def sample_job():
    return JobListing(
        job_id="abc123",
        title="Data Analyst",
        company="Example Corp",
        location="Remote",
        description="Analyze data with SQL and Python.",
        requirements="SQL, Python, Excel",
        application_url="https://example.com/jobs/1",
    )


def test_resume_retriever_falls_back_to_single_template(tmp_path):
    from agents.resume_retriever import ResumeRetriever

    template = tmp_path / "Resume.docx"
    template.write_text("fake docx")

    retriever = ResumeRetriever(tmp_path)
    selected = retriever.retrieve(
        JobListing(title="Data Analyst", company="Test", application_url="https://example.com"),
        profile={},
    )
    assert selected == template


def test_resume_retriever_picks_best_matching_template(tmp_path):
    from agents.resume_retriever import ResumeRetriever
    from docx import Document

    ai_template = tmp_path / "Resume AI Engineer.docx"
    analyst_template = tmp_path / "Resume Data Analyst.docx"
    generic_template = tmp_path / "Resume Generic.docx"

    for path in (ai_template, analyst_template, generic_template):
        doc = Document()
        doc.add_paragraph(path.stem)
        doc.save(str(path))

    retriever = ResumeRetriever(tmp_path)
    selected = retriever.retrieve(
        JobListing(
            title="Senior Data Analyst",
            company="Test",
            description="SQL Python Tableau dashboards",
            application_url="https://example.com",
        ),
        profile={"preferences": {"target_roles": ["Data Analyst"]}},
    )
    assert selected.name == "Resume Data Analyst.docx"


def test_resume_tailor_extracts_date_ranges(tmp_path):
    from agents.resume_tailor import ResumeTailor
    from docx import Document

    doc = Document()
    doc.add_paragraph("Data Analyst | PepsiCo | Dec 2023 – Present")
    doc.add_paragraph("Data Analyst | Coca-Cola | Jan 2020 – Nov 2023")
    template_path = tmp_path / "resume.docx"
    doc.save(str(template_path))

    tailor = ResumeTailor(template_path)
    assert any("Dec 2023" in dr for dr in tailor.date_ranges)
    assert any("Jan 2020" in dr for dr in tailor.date_ranges)


def test_resume_tailor_uses_fabrication_tolerance(tmp_path):
    from agents.resume_tailor import ResumeTailor
    from docx import Document

    doc = Document()
    doc.add_paragraph("Data Analyst | PepsiCo | Dec 2023 – Present")
    doc.add_paragraph("Built dashboards")
    template_path = tmp_path / "resume.docx"
    doc.save(str(template_path))

    tailor = ResumeTailor(template_path)
    job = JobListing(
        title="AI Engineer",
        company="OpenAI",
        description="Need LLM and Python skills.",
        application_url="https://example.com",
    )

    mock_analysis = {
        "role_title": "AI Engineer",
        "key_skills": ["Python", "LLMs"],
        "required_experience": "2 years",
        "soft_skills": ["communication"],
        "keywords_for_ats": ["Python", "LLM"],
        "resume_title": "AI Engineer | Python | LLM",
        "summary_focus": "Build AI systems",
        "top_achievements_to_highlight": ["AI projects"],
        "cover_letter_angle": "AI expertise",
    }

    mock_content = {
        "professional_title": "AI Engineer | Python",
        "professional_summary": "Summary",
        "technical_skills": [{"category": "Skills", "skills": "Python"}],
        "experience": [{"job_header": "Data Analyst | PepsiCo | Dec 2023 – Present", "bullets": ["Built AI features"]}],
    }

    with patch("agents.resume_tailor.JDAnalyzer") as mock_analyzer_cls:
        mock_analyzer_cls.return_value.analyze.return_value = mock_analysis
        with patch("agents.resume_tailor.llm_client") as mock_llm:
            mock_llm.chat.return_value = f"```json\n{json.dumps(mock_content)}\n```"
            content = tailor.tailor(job, fabrication_tolerance="aggressive")

    assert content["professional_title"] == "AI Engineer | Python"
    # The prompt should have mentioned aggressive tolerance.
    prompt = mock_llm.chat.call_args[0][0][1]["content"]
    assert "AGGRESSIVE" in prompt or "aggressive" in prompt


def test_ats_recruiter_scorer_parses_llm_output(sample_job):
    from agents.ats_recruiter_scorer import ATSRecruiterScorer

    scorer = ATSRecruiterScorer()
    with patch("agents.ats_recruiter_scorer.llm_client") as mock_llm:
        mock_llm.chat.return_value = '{"ats_score": 85, "recruiter_score": 80, "feedback": ""}'
        result = scorer.score({"professional_title": "Data Analyst"}, sample_job)

    assert result["ats_score"] == 85
    assert result["recruiter_score"] == 80
    assert result["passed"] is True
    assert result["feedback"] == ""


def test_ats_recruiter_scorer_fails_below_threshold(sample_job):
    from agents.ats_recruiter_scorer import ATSRecruiterScorer

    scorer = ATSRecruiterScorer(min_ats_score=80, min_recruiter_score=80)
    with patch("agents.ats_recruiter_scorer.llm_client") as mock_llm:
        mock_llm.chat.return_value = '{"ats_score": 60, "recruiter_score": 75, "feedback": "Add more keywords"}'
        result = scorer.score({"professional_title": "Data Analyst"}, sample_job)

    assert result["passed"] is False
    assert "Add more keywords" in result["feedback"]


def test_consistency_ledger_records_entry(tmp_path, sample_job):
    ledger_path = tmp_path / "ledger.json"
    ledger = ConsistencyLedger(ledger_path)

    content = {
        "professional_title": "Data Analyst | SQL",
        "professional_summary": "Summary",
        "technical_skills": [{"category": "Skills", "skills": "SQL"}],
        "experience": [{"job_header": "Analyst | Co | Jan 2020 – Present", "bullets": ["Did things"]}],
    }
    entry = ledger.record(
        job=sample_job,
        source_template=tmp_path / "template.docx",
        fabrication_tolerance="moderate",
        tailored_content=content,
        output_paths={"resume_pdf": tmp_path / "out.pdf"},
        revision=1,
    )

    assert entry["job_id"] == "abc123"
    assert entry["fabrication_tolerance"] == "moderate"
    assert entry["claims"]["professional_title"] == "Data Analyst | SQL"
    assert (tmp_path / "ledger.json").exists()

    entries = ledger.get_entries_for_job("abc123")
    assert len(entries) == 1


def test_consistency_ledger_returns_empty_for_missing_file(tmp_path):
    ledger = ConsistencyLedger(tmp_path / "does_not_exist.json")
    assert ledger.list_entries() == []
